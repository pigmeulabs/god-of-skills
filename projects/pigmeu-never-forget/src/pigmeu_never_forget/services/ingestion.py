"""Ingestion service with snapshot, diff and checkpoint persistence."""

from __future__ import annotations

import hashlib
import json
import re
from html import escape
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path

from pigmeu_never_forget.domain.models import ProjectContext, SourceDocument
from pigmeu_never_forget.storage.sqlite import sqlite_cursor

try:
    from markdownify import markdownify as md_to_markdown
except ModuleNotFoundError:  # pragma: no cover - safe fallback for minimal environments
    md_to_markdown = None

try:
    from pypdf import PdfReader
except ModuleNotFoundError:  # pragma: no cover - optional dependency
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ModuleNotFoundError:  # pragma: no cover - optional dependency
    DocxDocument = None

try:
    from PIL import Image
except ModuleNotFoundError:  # pragma: no cover - optional dependency
    Image = None

PARSER_VERSION = "v1"
SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".pdf",
    ".docx",
    ".html",
    ".json",
    ".csv",
    ".py",
    ".js",
    ".ts",
    ".png",
    ".jpg",
    ".jpeg",
    ".gif",
    ".webp",
    ".bmp",
    ".tiff",
    ".svg",
}


@dataclass(slots=True)
class DiscoveredFile:
    """A candidate file for ingestion."""

    path: Path
    relative_path: str
    source_hash: str


@dataclass(slots=True)
class SourceSnapshot:
    """Normalized snapshot of a source file for diffing."""

    path: str
    source_hash: str
    content_hash: str
    parser_version: str
    document: SourceDocument


@dataclass(slots=True)
class FileCheckpoint:
    """Persisted checkpoint row."""

    path: str
    project_id: str
    source_hash: str
    content_hash: str
    parser_version: str
    last_scan_at: str


@dataclass(slots=True)
class IngestionDiff:
    """Diff between current source snapshots and persisted checkpoints."""

    new: list[SourceSnapshot] = field(default_factory=list)
    modified: list[SourceSnapshot] = field(default_factory=list)
    unchanged: list[SourceSnapshot] = field(default_factory=list)
    deleted: list[FileCheckpoint] = field(default_factory=list)
    renamed: list[tuple[FileCheckpoint, SourceSnapshot]] = field(default_factory=list)

    def to_summary(self) -> dict[str, int]:
        """Return diff counters."""
        return {
            "new": len(self.new),
            "modified": len(self.modified),
            "unchanged": len(self.unchanged),
            "deleted": len(self.deleted),
            "renamed": len(self.renamed),
        }


class IngestionService:
    """Discover, normalize and diff project inputs."""

    def discover_files(self, project: ProjectContext) -> list[DiscoveredFile]:
        """Return eligible files according to project settings."""
        include_paths = project.settings["project_settings"]["indexing"]["include_paths"]
        exclude_patterns = set(project.settings["project_settings"]["indexing"]["exclude_patterns"])
        discovered: list[DiscoveredFile] = []
        for include_path in include_paths:
            target = project.root_path / include_path
            if not target.exists():
                continue
            for file_path in target.rglob("*"):
                if not file_path.is_file():
                    continue
                if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                    continue
                if any(part in exclude_patterns for part in file_path.parts):
                    continue
                discovered.append(
                    DiscoveredFile(
                        path=file_path,
                        relative_path=str(file_path.relative_to(project.root_path)),
                        source_hash=self.compute_source_hash(file_path),
                    )
                )
        return sorted(discovered, key=lambda item: item.relative_path)

    def normalize_file(self, project: ProjectContext, file_path: Path) -> SourceDocument:
        """Normalize a text file into the canonical source document shape."""
        text = self._read_and_normalize(file_path)
        relative_path = file_path.relative_to(project.root_path)
        return SourceDocument(
            source_id=str(relative_path).replace("/", "__"),
            source_type="file",
            path=str(relative_path),
            title=file_path.stem,
            mime_type=self._guess_mime_type(file_path),
            raw_text=text,
            metadata={"normalized": True},
        )

    def build_snapshots(self, project: ProjectContext) -> list[SourceSnapshot]:
        """Build normalized snapshots for the current file discovery set."""
        snapshots: list[SourceSnapshot] = []
        for item in self.discover_files(project):
            document = self.normalize_file(project, item.path)
            snapshots.append(
                SourceSnapshot(
                    path=item.relative_path,
                    source_hash=item.source_hash,
                    content_hash=self.compute_content_hash(document.raw_text),
                    parser_version=PARSER_VERSION,
                    document=document,
                )
            )
        return snapshots

    def diff_snapshots(self, project: ProjectContext, snapshots: list[SourceSnapshot]) -> IngestionDiff:
        """Compare current snapshots with persisted file checkpoints."""
        checkpoints = self._load_checkpoints(project)
        current_by_path = {snapshot.path: snapshot for snapshot in snapshots}
        diff = IngestionDiff()

        for path, snapshot in current_by_path.items():
            checkpoint = checkpoints.get(path)
            if checkpoint is None:
                diff.new.append(snapshot)
                continue
            if (
                checkpoint.source_hash == snapshot.source_hash
                and checkpoint.content_hash == snapshot.content_hash
                and checkpoint.parser_version == snapshot.parser_version
            ):
                diff.unchanged.append(snapshot)
            else:
                diff.modified.append(snapshot)

        for path, checkpoint in checkpoints.items():
            if path not in current_by_path:
                diff.deleted.append(checkpoint)

        self._mark_renames(diff)
        return diff

    def persist_checkpoints(
        self,
        project: ProjectContext,
        snapshots: list[SourceSnapshot],
        diff: IngestionDiff,
    ) -> None:
        """Persist latest file checkpoints after applying the diff."""
        now = datetime.now(tz=timezone.utc).isoformat()
        with sqlite_cursor(project.state_db_path) as connection:
            for snapshot in snapshots:
                connection.execute(
                    """
                    insert into file_checkpoints(path, project_id, source_hash, content_hash, parser_version, last_scan_at)
                    values (?, ?, ?, ?, ?, ?)
                    on conflict(path) do update set
                      project_id = excluded.project_id,
                      source_hash = excluded.source_hash,
                      content_hash = excluded.content_hash,
                      parser_version = excluded.parser_version,
                      last_scan_at = excluded.last_scan_at
                    """,
                    (
                        snapshot.path,
                        project.project_id,
                        snapshot.source_hash,
                        snapshot.content_hash,
                        snapshot.parser_version,
                        now,
                    ),
                )
            for old_checkpoint, _ in diff.renamed:
                connection.execute("delete from file_checkpoints where path = ?", (old_checkpoint.path,))
            for checkpoint in diff.deleted:
                connection.execute("delete from file_checkpoints where path = ?", (checkpoint.path,))

    def list_pending_vector_reconciliation(self, project: ProjectContext) -> list[str]:
        """Return active chunks that still miss vector confirmation."""
        with sqlite_cursor(project.state_db_path) as connection:
            rows = connection.execute(
                """
                select chunk_id
                from chunks
                where active = 1 and (qdrant_point_id is null or qdrant_point_id = '')
                order by chunk_id
                """
            ).fetchall()
        return [row["chunk_id"] for row in rows]

    def mark_chunk_vector_confirmed(self, project: ProjectContext, chunk_id: str, qdrant_point_id: str) -> None:
        """Store vector confirmation for a chunk after successful upsert."""
        with sqlite_cursor(project.state_db_path) as connection:
            connection.execute(
                "update chunks set qdrant_point_id = ? where chunk_id = ?",
                (qdrant_point_id, chunk_id),
            )

    def _load_checkpoints(self, project: ProjectContext) -> dict[str, FileCheckpoint]:
        with sqlite_cursor(project.state_db_path) as connection:
            rows = connection.execute(
                """
                select path, project_id, source_hash, content_hash, parser_version, last_scan_at
                from file_checkpoints
                where project_id = ?
                """,
                (project.project_id,),
            ).fetchall()
        return {row["path"]: FileCheckpoint(**dict(row)) for row in rows}

    def _mark_renames(self, diff: IngestionDiff) -> None:
        deleted_pool = list(diff.deleted)
        rename_pairs: list[tuple[FileCheckpoint, SourceSnapshot]] = []
        remaining_new: list[SourceSnapshot] = []
        for snapshot in diff.new:
            matched: FileCheckpoint | None = None
            for checkpoint in deleted_pool:
                if checkpoint.content_hash == snapshot.content_hash:
                    matched = checkpoint
                    break
            if matched is None:
                remaining_new.append(snapshot)
                continue
            rename_pairs.append((matched, snapshot))
            deleted_pool.remove(matched)
        diff.renamed = rename_pairs
        diff.new = remaining_new
        diff.deleted = deleted_pool

    def _read_and_normalize(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix == ".json":
            raw = file_path.read_text(encoding="utf-8", errors="replace")
            html = self._normalize_json_to_html(raw)
            return self._html_to_markdown(html)
        if suffix == ".csv":
            raw = file_path.read_text(encoding="utf-8", errors="replace")
            html = self._normalize_csv_to_html(raw)
            return self._html_to_markdown(html)
        if suffix == ".html":
            raw = file_path.read_text(encoding="utf-8", errors="replace")
            html = self._normalize_html(raw)
            return self._html_to_markdown(html)
        if suffix == ".pdf":
            html = self._extract_pdf_to_html(file_path)
            return self._html_to_markdown(html)
        if suffix == ".docx":
            html = self._extract_docx_to_html(file_path)
            return self._html_to_markdown(html)
        if suffix in {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".svg"}:
            html = self._extract_image_to_html(file_path)
            return self._html_to_markdown(html)
        raw = file_path.read_text(encoding="utf-8", errors="replace")
        html = f"<pre>{escape(raw)}</pre>"
        return self._html_to_markdown(html)

    def _normalize_json_to_html(self, raw: str) -> str:
        try:
            payload = json.loads(raw)
            pretty = json.dumps(payload, ensure_ascii=True, sort_keys=True, indent=2)
            return f"<h1>JSON</h1><pre>{escape(pretty)}</pre>"
        except json.JSONDecodeError:
            return f"<h1>JSON</h1><pre>{escape(raw)}</pre>"

    def _normalize_csv_to_html(self, raw: str) -> str:
        lines = [line.strip() for line in raw.splitlines() if line.strip()]
        if not lines:
            return "<h1>CSV</h1><p>Empty file</p>"
        rows = [line.split(",") for line in lines]
        header = rows[0]
        body = rows[1:]
        table_header = "".join(f"<th>{escape(cell.strip())}</th>" for cell in header)
        table_rows = "".join(
            "<tr>" + "".join(f"<td>{escape(cell.strip())}</td>" for cell in row) + "</tr>"
            for row in body
        )
        return f"<h1>CSV</h1><table><thead><tr>{table_header}</tr></thead><tbody>{table_rows}</tbody></table>"

    def _normalize_html(self, raw: str) -> str:
        no_scripts = re.sub(r"<script[\s\S]*?</script>", " ", raw, flags=re.IGNORECASE)
        no_styles = re.sub(r"<style[\s\S]*?</style>", " ", no_scripts, flags=re.IGNORECASE)
        cleaned = re.sub(r"\s+", " ", no_styles).strip()
        return cleaned

    def _binary_to_html_stub(self, file_path: Path, file_kind: str) -> str:
        return (
            f"<h1>{escape(file_kind.upper())} Document</h1>"
            f"<p>Source file: <code>{escape(file_path.name)}</code></p>"
            "<p>Content extraction parser will be attached in a later stage.</p>"
        )

    def _extract_pdf_to_html(self, file_path: Path) -> str:
        if PdfReader is None:
            return self._binary_to_html_stub(file_path, file_kind="pdf")
        try:
            reader = PdfReader(str(file_path))
            page_blocks: list[str] = []
            for index, page in enumerate(reader.pages, start=1):
                extracted = page.extract_text() or ""
                if extracted.strip():
                    page_blocks.append(
                        f"<h2>Page {index}</h2><pre>{escape(extracted.strip())}</pre>"
                    )
            if not page_blocks:
                return self._binary_to_html_stub(file_path, file_kind="pdf")
            return f"<h1>PDF Document</h1><p>Source file: <code>{escape(file_path.name)}</code></p>{''.join(page_blocks)}"
        except Exception:
            return self._binary_to_html_stub(file_path, file_kind="pdf")

    def _extract_docx_to_html(self, file_path: Path) -> str:
        if DocxDocument is None:
            return self._binary_to_html_stub(file_path, file_kind="docx")
        try:
            doc = DocxDocument(str(file_path))
            paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
            if not paragraphs:
                return self._binary_to_html_stub(file_path, file_kind="docx")
            body = "".join(f"<p>{escape(paragraph)}</p>" for paragraph in paragraphs)
            return f"<h1>DOCX Document</h1><p>Source file: <code>{escape(file_path.name)}</code></p>{body}"
        except Exception:
            return self._binary_to_html_stub(file_path, file_kind="docx")

    def _extract_image_to_html(self, file_path: Path) -> str:
        if Image is None:
            return self._binary_to_html_stub(file_path, file_kind="image")
        try:
            with Image.open(file_path) as image:
                width, height = image.size
                mode = image.mode
            return (
                "<h1>IMAGE Document</h1>"
                f"<p>Source file: <code>{escape(file_path.name)}</code></p>"
                f"<p>Dimensions: {width}x{height}</p>"
                f"<p>Color mode: {escape(mode)}</p>"
            )
        except Exception:
            return self._binary_to_html_stub(file_path, file_kind="image")

    def _html_to_markdown(self, html: str) -> str:
        if md_to_markdown is not None:
            return md_to_markdown(html, heading_style="ATX").strip()
        # Fallback for minimal environments where markdownify is not installed.
        text = re.sub(r"<[^>]+>", " ", html)
        return re.sub(r"\s+", " ", text).strip()

    @staticmethod
    def compute_source_hash(file_path: Path) -> str:
        """Hash original file bytes."""
        digest = hashlib.sha256()
        with file_path.open("rb") as handle:
            for chunk in iter(lambda: handle.read(8192), b""):
                digest.update(chunk)
        return digest.hexdigest()

    @staticmethod
    def compute_content_hash(text: str) -> str:
        """Hash normalized text content."""
        return hashlib.sha256(text.encode("utf-8")).hexdigest()

    @staticmethod
    def _guess_mime_type(file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        return {
            ".md": "text/markdown",
            ".txt": "text/plain",
            ".py": "text/x-python",
            ".json": "application/json",
            ".csv": "text/csv",
            ".html": "text/html",
        }.get(suffix, "application/octet-stream")
